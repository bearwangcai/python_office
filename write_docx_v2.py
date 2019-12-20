import os  
import docx
from docx import Document
from read_docx_v2 import read_file
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from copy import deepcopy

global professions
global province_names #省名
global content_names #共性意见 个性意见
global all



def cell_merge(table):
    cell_merge_list = []
    for row_index in range(len(table.rows)):
        for col_index in range(len(table.columns)):
            cell = table.cell(row_index,col_index)   #遍历所有单元格
            col_list = []   #将已经合并单元格的点放入列表中
            if (row_index,col_index) in cell_merge_list:   #已合并的单元格不再重复查验
                break
            else:
                try:
                    _ = float(cell.text)   #纯数字类单元格不合并
                except:
                    for sub_row_index in range(row_index,len(table.rows)):
                        if table.cell(sub_row_index,col_index).text == cell.text:
                            row = sub_row_index   #记录最大相同内容的行号，最小相同内容行号为本身行号
                            for sub_col_index in range(col_index,len(table.columns)):
                                if table.cell(sub_row_index,sub_col_index).text == cell.text:
                                    col_proxy = sub_col_index
                                    if col_proxy == (len(table.columns)-1):  #如果改行最右列内容相同，列号计入列表
                                        col_list.append(col_proxy)
                                else:
                                    col_list.append(col_proxy)   #记录内容相同的最右列号
                                    break   #如果同行下一轮内容不相同，直接不再向右搜索
                        else: 
                            break   #如果下一行同列内容不相同，直接不再向下搜索
                    if (row != row_index or min(col_list) != col_index):  #只要不是同一单元格，合并
                        text_cache = cell.text
                        cell.merge(table.cell(row, min(col_list)))
                        cell.text = text_cache  #内容只保留一个即可
                        for i in range(row+1):
                            for j in range(min(col_list)+1):
                                cell_merge_list.append((i,j))  #将已经合并单元格的点放入列表中


def write_data(document, content, line_spacing, space_after):
    if isinstance(content, Paragraph):
        if not (content.text == ''):
            
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.first_line_indent = Pt(32) #首行缩进

            for run_origin in content.runs:
                run = p.add_run(run_origin.text)
                run.bold = run_origin.bold
                run.italic = run_origin.italic
                run.font.color.rgb = run_origin.font.color.rgb
    elif isinstance(content, Table):
        columns_num = len(content.columns)
        rows_num = len(content.rows)
        table = document.add_table(rows = rows_num, cols = columns_num, style = 'Table Grid')
        for item_row, row in enumerate(content.rows):   # 读每行
            hdr_cells = table.rows[item_row].cells
            for item_cell,cell in enumerate(row.cells):  # 读一行中的所有单元格
                c = cell.text
                hdr_cells[item_cell].text = c
        cell_merge(table)


def write_content(document, file, profession, province_name, line_spacing, space_after):
    '''
    写入具体评审意见，包含共性问题，个性问题等内容
    '''
    for content_name in content_names: #共性问题，个性问题
        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
        p.paragraph_format.space_after = Pt(space_after)
        if content_name == '（一）共性意见':
            p.add_run('（一）你省共性要求').bold = True
            for content in file[profession][content_name]:
                write_data(document, content, line_spacing, space_after)
        else:
            p.add_run('（二）你省其他审查意见').bold = True
            for content in file[profession][province_name][content_name]:
                write_data(document, content, line_spacing, space_after)
                    

def set_style(run, font_name = u'黑体', font_size = 18, bold = True):
    '''
    run
    font_name = u'黑体'
    font_size = 18  
    bold = True
    '''
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) #更改字体必须.font.name和._element.rPr.rFonts.set两个都要改
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x0a,0x0a,0x0a)
    run.bold = bold

def write_file(file):

    for province_name in province_names:
        document = Document()
        document.styles['Normal'].font.name = u'仿宋'
        document.styles['Normal'].font.size = Pt(16)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        line_spacing = 32 #行间距
        space_after =  0 #段后间距
        
        #标题行
        p = document.add_paragraph('')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #只有add_paragraph方法有这个属性
        run = p.add_run(province_name + '分公司2020-2022年网络发展滚动规划') #只有add_paragraph.add_run这个方法有下面的属性，而且想要更改个别属性，就需要用.add_run方法！
        set_style(run, u'黑体', 18, True)

        p = document.add_paragraph('')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('初审评审意见')
        set_style(run, u'黑体', 18, True)
        
        
        p = document.add_paragraph('')
        p = document.add_paragraph('    2019年10月28日-10月30日集团公司网络发展部组织集团规划项目组对你省公司2020-2022年网络发展滚动规划总体思路进行了集中评审，形成审查意见如下。')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        p = document.add_paragraph('')


        #网络规划
        run = document.add_heading('',level=1).add_run('I 网络规划') 
        set_style(run, u'黑体', 18, False)

        # 一 总体要求
        run = document.add_heading('',level=2).add_run('一 总体要求')
        set_style(run, u'黑体', 16, True) 

        p = document.add_paragraph('    （一）细化梳理网络现状，充分利旧现网资源，提质增效。')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        p = document.add_paragraph('    （二）规划方案原则上必须有国家要求（当地政府要求）、集团战略、市场需求或技术发展趋势等作为规划输入。')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        p = document.add_paragraph('    （三）为保证2020年的5G建设投资，除5G以外的各专业投资预计将大幅压缩，')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run('具体见各专业压缩比例。')
        #run.font.color.rgb = RGBColor(0xff,0x00,0x00) #想要个别更改属性，就需要用.add_run方法
        run1 = p.add_run('可建可不建的项目原则上不予建设。')


        # 二 各专业具体要求  
        p = document.add_paragraph('')
        run = document.add_heading('',level=2).add_run('二 各专业具体要求')
        set_style(run, u'黑体', 16, True)     

        for index_1, profession in enumerate(professions):
            
            #document.add_heading(profession, level=1)

            try:
                if (len(profession) == 1):
                    profession = profession[0]

                    #添加大专业标题
                    run = document.add_heading('',level=3).add_run('2.' + str(index_1+1) + ' ' + profession)
                    set_style(run, u'宋体', 16, True)
                    run.font.color.rgb = RGBColor(0x1a,0x1a,0x1a)
                    write_content(document, file, profession, province_name, line_spacing, space_after)
                else:
                    for index_2, sub_profession in enumerate(profession):

                        if index_2 == 0:
                            #添加大专业标题
                            run = document.add_heading('',level=3).add_run('2.' + str(index_1+1) + ' ' + sub_profession)
                            set_style(run, u'宋体', 16, True)
                            run.font.color.rgb = RGBColor(0x1a,0x1a,0x1a)

                        else:
                            #添加小专业标题
                            run = document.add_heading('',level=4).add_run('2.' + str(index_1+1) + '.' + str(index_2) + ' ' + sub_profession)
                            set_style(run, u'宋体 (中文标题)', 16, True)
                            run.font.color.rgb = RGBColor(0x1a,0x1a,0x1a)
                            write_content(document, file, sub_profession, province_name, line_spacing, space_after)
            except: 
                print(profession)
        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
        p.paragraph_format.space_after = Pt(space_after)
        '''
        #节能减排
        p = document.add_paragraph('')
        run = p.add_run('节能减排')
        set_style(run, u'黑体', 18, False)
        write_content(document, file, '节能减排', province_name, line_spacing, space_after)
        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
        p.paragraph_format.space_after = Pt(space_after)
        
        #应急通信
        run = document.add_heading('',level=1).add_run('II 应急通信')
        set_style(run, u'黑体', 18, False)
        write_content(document, file, '应急通信', province_name, line_spacing, space_after)
        '''

        document.save(r'C:\Users\x270\Desktop\test' + '\\'+ province_name + '分公司2020-2022年网络发展滚动规划初审意见.docx')



def main():
    global professions
    professions = [['CDN']]
    
    #professions = [['数据网'],['STN(IPRAN)'],['传输网'],['接入网'],['IDC'],['DC'],['云'],['业务平台'],['CDN'],['移动核心网'],['固网核心网'],['应急通信'],['节能减排']]

    #professions = [['移动网','核心网'],['基础网','接入网和综合业务接入区','IP网','STN(IPRAN)','传输网'],['IDC及基础设施建设','IDC及DC基础设施建设','管道及其他基础设施'],['云计算','云','CDN','业务平台']]

    global province_names #省名
    province_names = ['北京','天津','河北','山西','内蒙古','辽宁','吉林','黑龙江','上海','江苏','浙江','安徽','福建','江西','山东','河南','湖北','湖南','广东','广西','海南','重庆','四川','贵州','云南','西藏','陕西','甘肃','青海','宁夏','新疆']
    global content_names #共性意见 个性意见
    content_names = ['（一）共性意见','（二）你省审查意见']
    global all
    all = read_file()
    write_file(all)

if __name__ == "__main__":
    main()




        
