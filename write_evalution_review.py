import os  
import docx
from docx import Document
from read_evalution_review import read
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from copy import deepcopy

def cell_merge(table):
    cell_merge_list = []

    for row_index in range(len(table.rows)):
        cell = table.cell(row_index,0)   #遍历第一列单元格
        if (row_index,0) in cell_merge_list:   #已合并的单元格不再重复查验
            continue
        else:
            for sub_row_index in range(row_index,len(table.rows)):
                if table.cell(sub_row_index,0).text == cell.text:
                    row = sub_row_index   #记录最大相同内容的行号，最小相同内容行号为本身行号
                else: 
                    break   #如果下一行同列内容不相同，直接不再向下搜索
            if (row != row_index):  #只要不是同一单元格，合并
                text_cache = cell.text
                cell.merge(table.cell(row, 0))
                cell.text = text_cache  #内容只保留一个即可
                for i in range(row+1):
                    cell_merge_list.append((i,0))  #将已经合并单元格的点放入列表中    

def write_table(document, data, province_name, professions):
    '''
    按省生成所有专业项目核准表
    '''


    table = document.add_table(rows = 1, cols = 8, style = 'Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目类别'
    hdr_cells[1].text = '序号'
    hdr_cells[2].text = '项目名称'
    hdr_cells[3].text = '建设内容及规模'
    hdr_cells[4].text = '投资（万元）'
    hdr_cells[5].text = '是否为核准项目'
    hdr_cells[6].text = '核准意见'
    hdr_cells[7].text = '集团审批可研'

    number = 1

    for profession in professions:   # 读每行
        try:
            content = data[province_name][profession]
            for items in content:
                row_cells = table.add_row().cells
                row_cells[0].text = profession
                #row_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for index, item in enumerate(items):
                    row_cells[index + 1].text = str(item)
                    #row_cells[index + 1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[index + 1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row_cells[1].text = str(number)
                row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER                
                number += 1
        except:
            pass
    cell_merge(table)


def set_style(run, font_name = u'黑体', font_size = 18, bold = True):
    '''
    run
    font_name = u'黑体'
    font_size = 18  
    bold = True
    '''
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) #更改字体必须.font.name和._element.rPr.rFonts.set两个都要改
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x0a,0x0a,0x0a)
    run.bold = bold



def write_doc(file_dir, province_names, professions, target_dir):
    data = read(file_dir)

    count = 0
    
    for province_name in province_names:
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal'].font.size = Pt(10)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        line_spacing = 32 #行间距
        space_after =  0 #段后间距

        
        #标题行
        p = document.add_paragraph('')
        #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #只有add_paragraph方法有这个属性
        run = p.add_run('附件2') #只有add_paragraph.add_run这个方法有下面的属性，而且想要更改个别属性，就需要用.add_run方法！
        set_style(run, u'黑体', 16, False)

        p = document.add_paragraph('')
        p = document.add_paragraph('')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #只有add_paragraph方法有这个属性
        run = p.add_run(province_name + '分公司2019年网络建设项目核准表') #只有add_paragraph.add_run这个方法有下面的属性，而且想要更改个别属性，就需要用.add_run方法！
        set_style(run, u'黑体', 18, True)

        p = document.add_paragraph('')

        write_table(document, data, province_name, professions)

        count += 1
        print(count)

        document.save(target_dir + '\\'+ province_name + '分公司2020-2022年网络发展滚动规划项目核准表.docx')


        

def main():
    professions = ['传输','CDN','IDC网络','STN','业务平台','云计算']
    province_names = ['北京','天津','河北','山西','内蒙古','辽宁','吉林','黑龙江','上海','江苏','浙江','安徽','福建','江西','山东','河南','湖北','湖南','广东','广西','海南','重庆','四川','贵州','云南','西藏','陕西','甘肃','青海','宁夏','新疆']
    file_dir = r"C:\Users\x270\Desktop\新建文件夹 (2)"
    target_dir = r'C:\Users\x270\Desktop\table_doc'
    write_doc(file_dir, province_names, professions, target_dir)

if __name__ == "__main__":
    main()